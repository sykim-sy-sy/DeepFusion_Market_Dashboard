from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# Main Title
title = doc.add_heading('Deepfusion AI 연구개발 추진계획 보고서', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
font.size = Pt(11)

sections = [
    ("1. 개요 및 목적", 
     "CES 2026 인공지능(AI) 부문 최고혁신상을 수상한 세계 유일의 4D 이미징 레이더 딥러닝(RAPA) 기술을 기반으로, 글로벌 자율주행 모빌리티 시장 선점 및 국방·해양 무인화 체계 수주를 가속하여 글로벌 탑티어 센서 퓨전(Sensor Fusion) 기업으로 도약하고자 함."),
    
    ("2. 추진방안", 
     "1) 글로벌 자율주행 시장 공략 (고가 센서 대체)\n - 고비용의 LiDAR 및 스테레오 비전을 완벽히 대체하는 4D 이미징 레이더 중심의 초기 융합(Early Fusion) 딥러닝 모델(RAPA-RC) 상용화.\n - BMW, Honda 등 글로벌 완성차(OEM) 및 DiDi, GAC 등 로보택시 운영사에 통합 제어기(Fusion ECU/SDK) 납품 추진.\n\n2) 국방 및 해양 무인화 사업 수주 확대\n - LIG Nex1 및 방위사업청과의 협력을 바탕으로 정찰·전투용 무인수상정(USV) 및 무인전투차량(UGV)에 특화된 근거리 전파계 통합 인지 시스템 공급 전력화.\n\n3) 글로벌 안전 규제 선제 대응 및 반복적 SW 매출 확보\n - 글로벌 안전 규제(Euro NCAP 2026, EU GSR2) 충족을 위한 딥러닝 기반 '에비던스 팩(Evidence Pack)' 모델 런칭.\n - 일회성 판매를 넘어선 지속적인 구독형 반복 매출 창출."),
    
    ("3. 실행방안", 
     "1) 하드웨어·소프트웨어 통합 레퍼런스 실증(PoC) 고도화 및 글로벌 고객사에 양산 라인(Design-in) 단계적 포팅 적용.\n\n2) 유럽 연구소(현지 거점) 설립을 통한 현지 고객 밀착 지원 체계 구축 및 정확도 개선을 위한 3만 건(30K) 분량의 4D 레이더 실차 데이터셋 조기 확보 달성."),
    
    ("4. 사업예산 규모 (총 50억원 가정)", 
     "1) 글로벌 R&D 연구 인력 확충 및 유럽 연구소 운영 (인건비/운영비) : 25억원 규모\n2) B2B 상용화용 4D 인지 솔루션 및 신규 융합 알고리즘 (RAPA-RL/RC) 고도화 (R&D 비용) : 15억원 규모\n3) 융합 제어기(Fusion ECU) 양산 조립 검사 라인 설치 및 실환경 실차 검증 인프라 구축 : 10억원 규모"),
    
    ("5. 기타사항 및 협조사항", 
     "1) 글로벌 실증(PoC) 인프라 지원: 해외 완성차 및 로보택시 업체와의 파일럿 테스트 시, 테스트 차량 배차 및 현지 주행 로깅 인프라 전폭적 지원 요망.\n\n2) 국방 파트너사(LIG Nex1 등) 연계 시험평가 협조: 무인수상정(USV) 체계 해상 실증을 위한 보안 인허가 및 행정 처리 최우선 지원 바람.\n\n3) 해외 진출 안전 기준 컴플라이언스 선제 검토: 향후 'Evidence Pack'의 북미/유럽 진출 시 필수적인 Euro NCAP 및 ISO 인증을 대비한 법무 파트의 사전 검토 협조 요망.")
]

for title_text, body_text in sections:
    # Heading (Level 1)
    h = doc.add_heading(title_text, level=1)
    # Ensure heading uses Malgun Gothic for Korean
    for r in h.runs:
        r.font.name = '맑은 고딕'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    # Body paragraph
    p = doc.add_paragraph()
    # Add spacing after paragraph
    p.paragraph_format.space_after = Pt(14)
    p.paragraph_format.line_spacing = 1.3
    
    run = p.add_run(body_text)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.font.size = Pt(11)

# save exactly to the path
doc.save(r'C:\Users\KIM SIYE\OneDrive - 딥퓨전에이아이주식회사\바탕 화면\SIYE\Marketing_Team\just\Deepfusion_깔끔한보고서.docx')
print("Successfully created the docx.")
